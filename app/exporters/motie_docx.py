from __future__ import annotations
from io import BytesIO
from pathlib import Path
from typing import Iterable, Dict, Any, Optional
from docx.shared import Pt, Cm
from docxtpl import DocxTemplate, InlineImage
from slugify import slugify
from flask import current_app, g
from docx.enum.style import WD_STYLE_TYPE
from collections import OrderedDict
import os



# ---- helpers ---------------------------------------------------------------
PREFERRED_BULLET_STYLE = "MotieBullet"  # maak deze in je DOCX (zie uitleg onderaan)

_BULLET_STYLE_CANDIDATES = [
    PREFERRED_BULLET_STYLE,
    "List Paragraph",                # EN
    "Bulleted List",                 # EN (soms)
    "List Bullet",                   # EN (soms)
    "Lijst met opsommingstekens",    # NL
    "Opsomming",                     # NL (soms)
    "Liste à puces",                 # FR
    "Aufzählungszeichen",            # DE
    "Párrafo de lista",              # ES
]

try:
    import requests  # alleen nodig voor externe logo_url
except Exception:
    requests = None


def _party_logo_inline(doc, party, width_cm=2.5):
    """Maak een InlineImage voor het partijlogo (lokaal bestand of via URL)."""
    if not party:
        return None

    # 1) lokaal bestand proberen
    if getattr(party, "logo_filename", None):
        local_path = os.path.join(current_app.static_folder, "img", "partijen", party.logo_filename)
        if os.path.exists(local_path):
            return InlineImage(doc, local_path, width=Cm(width_cm))

    # 2) URL proberen (optioneel)
    if getattr(party, "logo_url", None) and requests:
        try:
            r = requests.get(party.logo_url, timeout=6)
            r.raise_for_status()
            return InlineImage(doc, BytesIO(r.content), width=Cm(width_cm))
        except Exception:
            pass

    return None


def _chunk(lst, n):
    return [lst[i:i + n] for i in range(0, len(lst), n)]


def _find_available_style(doc: DocxTemplate) -> str | None:
    styles = doc.docx.styles
    names = {s.name for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH}
    for name in _BULLET_STYLE_CANDIDATES:
        if name in names:
            return name
    return None

def build_bullet_list(doc: DocxTemplate, items):
    sd = doc.new_subdoc()
    style_name = _find_available_style(doc)

    for txt in (items or []):
        t = str(txt).strip()
        if not t:
            continue

        if style_name:
            p = sd.add_paragraph(t, style=style_name)
        else:
            # Fallback: visueel identiek aan bullets (geen echte Word-lijst, maar strak en compact)
            p = sd.add_paragraph()
            p.add_run("• ")
            p.add_run(t)

        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1  # Enkel

    return sd
def _ensure_iter(value) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        # splits op newline of ; als je tekstveld hebt
        parts = [p.strip() for p in value.split("\n") if p.strip()]
        return parts
    if isinstance(value, Iterable):
        return [str(v).strip() for v in value if str(v).strip()]
    return [str(value)]

def _filename_for_motie(title: str) -> str:
    # 'Motie <titel>.docx' met veilige bestandsnaam
    safe = slugify(title or "", lowercase=False, separator=" ")
    safe = safe.strip()
    if not safe:
        safe = "Zonder titel"
    return f"Motie {safe}.docx"

def _template_path() -> Path:
    # Standaardlocatie voor je vaste template
    base = Path(current_app.root_path).parent  # project root /app/..
    # Tenant-specifieke template eerst proberen: app/templates_word/<tenant_slug>/motie.docx
    tenant = getattr(g, 'tenant', None)
    if tenant and getattr(tenant, 'slug', None):
        tp = base / "app" / "templates_word" / tenant.slug / "motie.docx"
        if tp.exists():
            return tp
    # Fallback naar algemene template
    return base / "app" / "templates_word" / "motie.docx"

# ---- public API ------------------------------------------------------------

def render_motie_to_docx_bytes(motie, *, vergadering: Optional[str] = None,
                               datum: Optional[str] = None) -> tuple[bytes, str]:
    tpl_path = _template_path()
    if not tpl_path.exists():
        raise FileNotFoundError(
            f"Word-sjabloon niet gevonden op {tpl_path}. "
            f"Zet je vaste format in app/templates_word/motie_template.docx of update _template_path()."
        )
    doc = DocxTemplate(str(tpl_path))

    # --- Basisvelden ---
    titel = getattr(motie, "titel", "") or ""
    opdracht_formulering = getattr(motie, "opdracht_formulering", "") or ""
    gemeenteraad_datum = getattr(motie, "gemeenteraad_datum", "") or ""
    agendapunt = getattr(motie, "agendapunt", "") or ""

    constaterende_dat = build_bullet_list(doc, _ensure_iter(getattr(motie, "constaterende_dat", [])))
    overwegende_dat   = build_bullet_list(doc, _ensure_iter(getattr(motie, "overwegende_dat", [])))
    draagt_college_op = build_bullet_list(doc, _ensure_iter(getattr(motie, "draagt_college_op", [])))

    # --- Ondertekenaars (hoofd + mede) ---
    ondertekenaars = []
    hoofd = getattr(motie, "indiener", None)
    if hoofd:
        ondertekenaars.append({
            "naam": hoofd.naam,
            "partij": getattr(hoofd.partij, "naam", None),
            "afkorting": getattr(hoofd.partij, "afkorting", None),
        })

    mede = list(getattr(motie, "mede_indieners", [])) or []
    mede_sorted = sorted(mede, key=lambda u: u.naam.casefold())
    for u in mede_sorted:
        ondertekenaars.append({
            "naam": u.naam,
            "partij": getattr(u.partij, "naam", None),
            "afkorting": getattr(u.partij, "afkorting", None),
        })

    # --- Unieke partijen (logo’s) in volgorde: hoofd eerst, dan mede ---
    unieke_partijen = OrderedDict()
    if hoofd and getattr(hoofd, "partij", None):
        unieke_partijen[hoofd.partij.id] = hoofd.partij
    for u in mede_sorted:
        if getattr(u, "partij", None):
            unieke_partijen[u.partij.id] = u.partij

    # Maak InlineImages
    partij_logo_items = []
    for p in unieke_partijen.values():
        partij_logo_items.append({
            "logo": _party_logo_inline(doc, p, width_cm=2.5),
            "naam": p.naam,
            "afkorting": p.afkorting,
        })

    # Optioneel: 6 per rij voor in een DOCX-tabel
    partij_logo_rows = _chunk(partij_logo_items, 6)

    context: Dict[str, Any] = {
        "titel": titel,
        "gemeenteraad_datum": gemeenteraad_datum,
        "agendapunt": agendapunt,
        "opdracht_formulering": opdracht_formulering,
        "constaterende_dat": constaterende_dat,
        "overwegende_dat": overwegende_dat,
        "draagt_college_op": draagt_college_op,

        # Nieuw:
        "ondertekenaars": ondertekenaars,             # lijst van dicts
        "partij_logo_rows": partij_logo_rows,         # lijst van rijen; elke rij is lijst van dicts met 'logo'
    }

    ondertekenaar_rows = _chunk(ondertekenaars, 6)

    context.update({
        "ondertekenaar_rows": ondertekenaar_rows,
    })


    doc.render(context)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read(), _filename_for_motie(titel)
